import random
import os
import requests
from docx import Document
from docx.shared import Inches, Pt

def create_word_with_random_color_image():
    # 1. רשימת צבעים לבחירה אקראית
    colors = ["red", "green", "blue", "yellow", "orange", "purple", "pink", "cyan"]
    selected_color = random.choice(colors)
    print(self_target := f"הצבע שנבחר אקראית: {selected_color}")

    # 2. הגדרת כתובת URL להורדת תמונה אקראית לפי הצבע שנבחר
    # נשתמש ב-LoremFlickr כדי לקבל תמונה רלוונטית לצבע
    image_url = f"https://loremflickr.com/800/600/{selected_color}"
    temp_image_path = "temp_color_image.jpg"

    print("מוריד תמונה מהרשת...")
    try:
        response = requests.get(image_url, timeout=15)
        if response.status_code == 200:
            with open(temp_image_path, 'wb') as f:
                f.write(response.content)
        else:
            print(f"שגיאה בהורדת התמונה: קוד סטטוס {response.status_code}")
            return
    except Exception as e:
        print(f"שגיאה בחיבור לרשת: {e}")
        return

    # 3. יצירת קובץ ה-Word והוספת התמונה
    print("מייצר קובץ Word...")
    doc = Document()
    
    # הוספת כותרת למסמך
    title = doc.add_paragraph()
    run = title.add_run(f"מסמך בהשראת הצבע: {selected_color.capitalize()}")
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph("להלן תמונה אקראית שנבחרה מהרשת בהתאם לצבע שנבחר:")

    # הוספת התמונה למסמך (ברוחב 5 אינץ')
    doc.add_picture(temp_image_path, width=Inches(5))

    # שמירת הקובץ
    output_filename = f"document_{selected_color}.docx"
    doc.save(output_filename)
    print(f"הקובץ נשמר בהצלחה בשם: {output_filename}")

    # 4. ניקוי קובץ התמונה הזמני מהמחשב
    if os.path.exists(temp_image_path):
        os.remove(temp_image_path)

if __name__ == "__main__":
    create_word_with_random_color_image()